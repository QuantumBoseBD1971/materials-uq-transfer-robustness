"""Build the anonymous journal-style DOCX manuscript from Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript_draft.md"
OUTPUT = ROOT / "materials_uq_frozen_threshold_manuscript_draft.docx"


def set_font(run, name="Times New Roman", size=11, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, text, separate, display, end):
        run._r.append(element)
    set_font(run, size=9)


def add_inline(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = paragraph.add_run(part.replace("`", ""))
            set_font(run)


def add_caption(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    set_font(run, size=9, italic=True)


def add_figure(document, filename, caption):
    path = ROOT / "figures" / filename
    if not path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(6.35))
    add_caption(document, caption)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after in (
        ("Heading 1", 14, 16, 8),
        ("Heading 2", 12, 12, 6),
        ("Heading 3", 11, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Number", "List Bullet"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15


def build():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.text = "Anonymous manuscript draft"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=9, italic=True)
        run.font.color.rgb = RGBColor(89, 89, 89)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(footer, "PAGE")

    lines = SOURCE.read_text().splitlines()
    inserted = set()
    in_equation = False
    equation_buffer = []
    for line in lines:
        stripped = line.strip()
        if stripped == "\\[":
            in_equation = True
            equation_buffer = []
            continue
        if stripped == "\\]":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(" ".join(equation_buffer))
            set_font(run, name="Cambria Math", size=11, italic=True)
            in_equation = False
            continue
        if in_equation:
            equation_buffer.append(stripped)
            continue
        if not stripped:
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(10)
            run = paragraph.add_run(stripped[2:])
            set_font(run, size=18, bold=True)
            continue
        if stripped.startswith("## "):
            heading = stripped[3:]
            document.add_heading(heading, level=1)
            if heading.startswith("4.2") and "operational" not in inserted:
                add_figure(
                    document,
                    "phase_d_operational_failures.png",
                    "Figure 1. Operational failure rates at the primary error targets. "
                    "Failure denotes a retained-risk breach or transferred coverage below 20%.",
                )
                inserted.add("operational")
            if heading.startswith("4.4") and "severity" not in inserted:
                add_figure(
                    document,
                    "phase_d_shift_severity.png",
                    "Figure 2. Retained-risk violation versus composition-distance shift severity. "
                    "The horizontal line marks the operating target.",
                )
                inserted.add("severity")
            if heading.startswith("4.5") and "tolerance" not in inserted:
                add_figure(
                    document,
                    "phase_d_tolerance_sensitivity.png",
                    "Figure 3. Operational failure rate across five error tolerances per task.",
                )
                inserted.add("tolerance")
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
            continue

        numbered = re.match(r"^\d+\.\s+(.*)", stripped)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(1))
            continue
        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            continue

        paragraph = document.add_paragraph()
        if stripped.startswith("**Anonymous manuscript"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(paragraph, stripped)

    document.core_properties.title = (
        "Evaluating the Transfer of Frozen Uncertainty-Based Error Controls "
        "Across Materials Chemistry Shifts"
    )
    document.core_properties.subject = "Anonymous manuscript draft"
    document.core_properties.author = "Independent Researcher"
    document.core_properties.keywords = (
        "materials informatics, uncertainty quantification, selective regression"
    )
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
